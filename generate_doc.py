from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ──
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2.5)

# ── Helper functions ──
def heading(text, level=1, color=RGBColor(0x1a, 0x56, 0xdb)):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = color
        run.font.bold = True
    return p

def body(text):
    p = doc.add_paragraph(text)
    p.style.font.size = Pt(11)
    return p

def bullet(text, bold_part=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.3)
    if bold_part and text.startswith(bold_part):
        run1 = p.add_run(bold_part)
        run1.bold = True
        run1.font.size = Pt(11)
        rest = text[len(bold_part):]
        run2 = p.add_run(rest)
        run2.font.size = Pt(11)
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.4)
    p.paragraph_format.right_indent = Inches(0.4)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F3F4F6')
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x1f, 0x29, 0x37)
    return p

def divider():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1A56DB')
        tcPr.append(shd)
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1].cells
        for ci, cell_text in enumerate(row_data):
            row[ci].text = cell_text
            for para in row[ci].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
        if ri % 2 == 0:
            for ci in range(len(row_data)):
                tc = table.rows[ri + 1].cells[ci]._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'EFF6FF')
                tcPr.append(shd)
    return table

# ════════════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════════════
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

title = title_para.add_run('FileVault')
title.font.size = Pt(36)
title.font.bold = True
title.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)

subtitle_para = doc.add_paragraph()
subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = subtitle_para.add_run('Project Documentation')
sub.font.size = Pt(18)
sub.font.color.rgb = RGBColor(0x6b, 0x72, 0x80)

doc.add_paragraph()
desc_para = doc.add_paragraph()
desc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
d = desc_para.add_run('Secure Cloud File Storage with Real-Time Monitoring')
d.font.size = Pt(13)
d.font.italic = True
d.font.color.rgb = RGBColor(0x37, 0x41, 0x51)

doc.add_paragraph()
doc.add_paragraph()

date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
dt = date_para.add_run(f'Date: {datetime.date.today().strftime("%B %d, %Y")}')
dt.font.size = Pt(11)
dt.font.color.rgb = RGBColor(0x9c, 0xa3, 0xaf)

doc.add_page_break()

# ════════════════════════════════════════════════
# 1. PROJECT OVERVIEW
# ════════════════════════════════════════════════
heading('1. Project Overview', 1)
body(
    'FileVault is a secure, web-based cloud file storage application built with Python and Flask. '
    'It allows users to upload, store, manage, and share files through a modern web interface, '
    'with enterprise-grade security features including multi-factor authentication, end-to-end '
    'encryption at rest, real-time activity monitoring, and biometric login support.'
)
doc.add_paragraph()
body('The application is composed of three main layers:')
bullet('Frontend: HTML/CSS/JavaScript with Bootstrap 5 for a responsive user interface.')
bullet('Backend: Python (Flask) handles all logic — authentication, file management, sharing, and security.')
bullet('Database: SQLite stores user accounts, file metadata, sharing records, and activity logs.')
divider()

# ════════════════════════════════════════════════
# 2. KEY FEATURES
# ════════════════════════════════════════════════
heading('2. Key Features', 1)

heading('2.1 User Authentication & Security', 2)
bullet('Secure registration and login with hashed passwords (bcrypt)')
bullet('Email verification — users must confirm their email address before full access')
bullet('Two-Factor Authentication (TOTP) — compatible with Google Authenticator and Authy')
bullet('WebAuthn / Biometric login — fingerprint or Face ID support on supported devices')
bullet('Account lockout after repeated failed login attempts')
bullet('Rate limiting to prevent brute-force attacks')
bullet('CSRF protection on all forms')
bullet('Content Security Policy (CSP) headers for XSS protection')

heading('2.2 File Management', 2)
bullet('Upload files (images, PDF, Word, Excel, PowerPoint, CSV, JSON, ZIP, MP4, MP3, and more)')
bullet('View, rename, and delete files')
bullet('File type filtering and sorting (by name, date, or size)')
bullet('Paginated file listing for large collections')
bullet('Drag-and-drop upload with real-time progress bar')

heading('2.3 File Sharing', 2)
bullet('Share files with specific registered users')
bullet('Generate public share links with optional expiry dates (1 day, 7 days, 30 days, or never)')
bullet('Revoke public links at any time')

heading('2.4 Real-Time Monitoring', 2)
bullet('Live activity notifications via WebSockets (Socket.IO)')
bullet('Activity log showing all uploads, deletions, and shares with timestamps')
bullet('Admin panel to manage all users and their files')

heading('2.5 User Interface', 2)
bullet('Responsive design — works on desktop, tablet, and mobile phones')
bullet('Dark mode and light mode toggle (preference saved per device)')
bullet('Collapsible sidebar navigation on small screens')
divider()

# ════════════════════════════════════════════════
# 3. TECHNOLOGY STACK
# ════════════════════════════════════════════════
heading('3. Technology Stack', 1)

add_table(
    ['Component', 'Technology', 'Purpose'],
    [
        ['Backend Framework', 'Flask 3.0 (Python)', 'Handles HTTP requests, routing, and business logic'],
        ['Database ORM', 'SQLAlchemy + Flask-Migrate', 'Manages database models and migrations'],
        ['Authentication', 'Flask-Login + Werkzeug', 'Session management and password hashing'],
        ['Two-Factor Auth', 'PyOTP + QRCode', 'TOTP code generation and QR code display'],
        ['Biometric Auth', 'WebAuthn (py library)', 'Passkey / fingerprint login support'],
        ['Real-Time', 'Flask-SocketIO + Eventlet', 'WebSocket connections for live notifications'],
        ['Email', 'Flask-Mail', 'Sends verification and password-reset emails'],
        ['Encryption', 'Cryptography (Fernet)', 'Encrypts uploaded files at rest'],
        ['Rate Limiting', 'Flask-Limiter', 'Prevents abuse and brute-force attacks'],
        ['Frontend', 'Bootstrap 5 + Bootstrap Icons', 'Responsive UI components and icons'],
        ['WSGI Server', 'Gunicorn + Eventlet worker', 'Production-grade web server'],
        ['Version Control', 'Git + GitHub', 'Code storage and deployment trigger'],
        ['Cloud Hosting', 'Render.com', 'Live production server'],
    ]
)
doc.add_paragraph()
divider()

# ════════════════════════════════════════════════
# 4. ROLE OF GITHUB
# ════════════════════════════════════════════════
heading('4. The Role of GitHub in This Project', 1)

heading('4.1 What is GitHub?', 2)
body(
    'GitHub is a cloud-based platform for storing and managing code using Git version control. '
    'Think of it as a secure online backup for your entire project, with a full history of every '
    'change ever made.'
)

heading('4.2 How GitHub is Used Here', 2)
bullet('Code Storage:', 'Code Storage:')
body(
    '     All the project files (Python code, HTML templates, CSS styles, JavaScript) are stored '
    'in a private repository named FileVault under the GitHub account Dahirou-Bachar. '
    'Private means only you can see and access it.'
)
doc.add_paragraph()
bullet('Version History:', 'Version History:')
body(
    '     Every time a change is made and committed, GitHub records exactly what changed, when, '
    'and why. You can always go back to any previous version of the code.'
)
doc.add_paragraph()
bullet('Deployment Trigger:', 'Deployment Trigger:')
body(
    '     Render.com is connected to this GitHub repository. Every time new code is pushed to '
    'GitHub, Render automatically detects the change and re-deploys the updated application '
    'within 2–3 minutes — no manual steps needed.'
)
doc.add_paragraph()
bullet('Disaster Recovery:', 'Disaster Recovery:')
body(
    '     If your computer is lost, broken, or reset, the entire project can be downloaded again '
    'from GitHub with a single command on any new device.'
)

heading('4.3 GitHub Repository Details', 2)
add_table(
    ['Property', 'Value'],
    [
        ['Repository Name', 'FileVault'],
        ['Account', 'Dahirou-Bachar'],
        ['Visibility', 'Private (only you can access it)'],
        ['URL', 'https://github.com/Dahirou-Bachar/FileVault'],
        ['Branch', 'main'],
    ]
)
doc.add_paragraph()
divider()

# ════════════════════════════════════════════════
# 5. ROLE OF RENDER
# ════════════════════════════════════════════════
heading('5. The Role of Render.com in This Project', 1)

heading('5.1 What is Render?', 2)
body(
    'Render.com is a cloud hosting platform that runs your application on a server 24 hours a day, '
    '7 days a week. Without a hosting service like Render, your app would only work when your '
    'personal computer is turned on and running.'
)

heading('5.2 How Render is Used Here', 2)
bullet('Web Server:', 'Web Server:')
body(
    '     Render runs the FileVault application using Gunicorn (a production-grade Python server) '
    'with the Eventlet worker for real-time WebSocket support. The exact start command is:'
)
code_block('gunicorn --worker-class eventlet -w 1 --bind 0.0.0.0:$PORT app:app')
doc.add_paragraph()
bullet('Always Online:', 'Always Online:')
body(
    '     The application is accessible at its public URL at all times. Users can log in, '
    'upload files, and use all features without your computer needing to be on.'
)
doc.add_paragraph()
bullet('Environment Variables:', 'Environment Variables:')
body(
    '     Sensitive configuration (secret keys, encryption keys, email passwords, WebAuthn settings) '
    'is stored securely in Render\'s environment variable settings — never in the code itself.'
)
doc.add_paragraph()
bullet('Auto-Deploy:', 'Auto-Deploy:')
body(
    '     Render is connected to the GitHub repository. When you push new code to GitHub, '
    'Render automatically pulls the latest version and restarts the application.'
)

heading('5.3 Render Service Details', 2)
add_table(
    ['Property', 'Value'],
    [
        ['Service Name', 'FileVault'],
        ['Service Type', 'Web Service (Free Tier)'],
        ['Live URL', 'https://filevault-eq7o.onrender.com'],
        ['Region', 'Oregon, USA'],
        ['Python Version', '3.11'],
        ['Start Command', 'gunicorn --worker-class eventlet -w 1 --bind 0.0.0.0:$PORT app:app'],
        ['Auto-Deploy', 'Yes — triggered on every GitHub push'],
    ]
)
doc.add_paragraph()

heading('5.4 Free Tier Limitation', 2)
body(
    'On Render\'s free plan, the server "sleeps" after 15 minutes of inactivity to save resources. '
    'The first request after sleep takes about 30–60 seconds to wake up — this is normal. '
    'Once awake, the app runs at full speed. Upgrading to a paid Render plan ($7/month) '
    'eliminates this sleep behavior.'
)
divider()

# ════════════════════════════════════════════════
# 6. HOW UPDATES WORK
# ════════════════════════════════════════════════
heading('6. How to Update the Project', 1)

body(
    'Because GitHub and Render are connected, updating the live application is a simple '
    '3-step process: edit the code, save it to GitHub, and Render does the rest automatically.'
)
doc.add_paragraph()

heading('6.1 The Update Workflow', 2)

add_table(
    ['Step', 'Action', 'Where'],
    [
        ['1', 'Make changes to the code on your computer', 'VS Code / any editor'],
        ['2', 'Save the changes to GitHub (git push)', 'Terminal / Command Prompt'],
        ['3', 'Render detects the push and re-deploys automatically', 'Render.com (automatic)'],
        ['4', 'Live app is updated in ~2–3 minutes', 'https://filevault-eq7o.onrender.com'],
    ]
)
doc.add_paragraph()

heading('6.2 Step-by-Step Commands to Push an Update', 2)
body('Open a terminal in your project folder and run these commands:')
doc.add_paragraph()

code_block('git add .')
body('     Stages all changed files, preparing them to be saved.')
doc.add_paragraph()

code_block('git commit -m "Describe what you changed here"')
body('     Saves the changes locally with a description message.')
doc.add_paragraph()

code_block('git push')
body('     Sends the changes to GitHub, which then triggers Render to re-deploy.')
doc.add_paragraph()

heading('6.3 What Happens During Render Re-Deployment', 2)
bullet('Render pulls the latest code from GitHub')
bullet('It installs any new Python packages listed in requirements.txt')
bullet('It restarts the Gunicorn server with the new code')
bullet('The live URL is updated — the whole process takes 2–3 minutes')
bullet('You can watch the build log in real time on the Render dashboard')

heading('6.4 Recovering and Continuing on a New Device', 2)
body('If you get a new computer or lose your current one, follow these steps:')
doc.add_paragraph()

body('Step 1 — Install the required software:')
bullet('Python 3.11 from python.org')
bullet('Git from git-scm.com')
bullet('VS Code (optional but recommended)')
doc.add_paragraph()

body('Step 2 — Download your project from GitHub:')
code_block('git clone https://github.com/Dahirou-Bachar/FileVault.git\ncd FileVault')
doc.add_paragraph()

body('Step 3 — Set up the Python environment:')
code_block('python -m venv venv\nvenv\\Scripts\\activate\npip install -r requirements.txt')
doc.add_paragraph()

body('Step 4 — Restore your secret keys:')
body(
    '     Your environment variables (SECRET_KEY, ENCRYPTION_KEY, MAIL_PASSWORD, etc.) '
    'are safely stored on Render. Go to:'
)
code_block('https://dashboard.render.com/web/srv-d81712gg4nts7398j0lg/env')
body('     Copy each value and create a .env file in the project folder.')
doc.add_paragraph()

body('Step 5 — Run the app locally:')
code_block('flask db upgrade\nflask run')
divider()

# ════════════════════════════════════════════════
# 7. ENVIRONMENT VARIABLES
# ════════════════════════════════════════════════
heading('7. Environment Variables (Secret Configuration)', 1)
body(
    'Environment variables store sensitive information that must never be written directly '
    'in the code. They are configured in Render\'s dashboard and loaded automatically when '
    'the application starts.'
)
doc.add_paragraph()

add_table(
    ['Variable', 'Purpose'],
    [
        ['SECRET_KEY', 'Encrypts user sessions and cookies'],
        ['ENCRYPTION_KEY', 'Encrypts uploaded files at rest using Fernet symmetric encryption'],
        ['MAIL_USERNAME', 'Gmail address used to send verification and reset emails'],
        ['MAIL_PASSWORD', 'Gmail App Password (not your regular password)'],
        ['WEBAUTHN_RP_ID', 'The domain name for biometric (WebAuthn) login — must match the live URL'],
        ['WEBAUTHN_RP_ORIGIN', 'The full HTTPS URL for biometric login — must match the live URL'],
        ['PYTHON_VERSION', 'Forces Render to use Python 3.11 for compatibility'],
    ]
)
doc.add_paragraph()

body(
    'IMPORTANT: Never share these values publicly. Never commit a .env file to GitHub. '
    'The .gitignore file in this project already excludes .env files from being uploaded.'
)
divider()

# ════════════════════════════════════════════════
# 8. PROJECT FILE STRUCTURE
# ════════════════════════════════════════════════
heading('8. Project File Structure', 1)
body('The main files and folders in the project:')
doc.add_paragraph()

add_table(
    ['File / Folder', 'Purpose'],
    [
        ['app.py', 'Main application file — all routes, logic, and configuration'],
        ['requirements.txt', 'Lists all Python packages the project depends on'],
        ['.python-version', 'Tells Render and local tools to use Python 3.11'],
        ['.gitignore', 'Lists files that should NOT be uploaded to GitHub (secrets, database, uploads)'],
        ['templates/', 'HTML pages (dashboard, login, register, profile, admin, etc.)'],
        ['static/css/', 'Style sheets — controls the look and feel of the app'],
        ['static/js/', 'JavaScript files — handles interactive behaviour (sidebar, uploads, etc.)'],
        ['instance/', 'Contains the SQLite database file (excluded from GitHub)'],
        ['vault_uploads/', 'Stores encrypted uploaded files (excluded from GitHub)'],
        ['migrations/', 'Database migration scripts managed by Flask-Migrate'],
    ]
)
doc.add_paragraph()
divider()

# ════════════════════════════════════════════════
# 9. SECURITY SUMMARY
# ════════════════════════════════════════════════
heading('9. Security Summary', 1)

add_table(
    ['Security Feature', 'How It Works'],
    [
        ['Password Hashing', 'Passwords are never stored as plain text — only an irreversible hash'],
        ['File Encryption', 'Every uploaded file is encrypted before being saved to disk'],
        ['TOTP 2FA', 'A 6-digit one-time code from an authenticator app is required at login'],
        ['WebAuthn / Biometrics', 'Fingerprint or Face ID login using the device\'s security chip'],
        ['CSRF Tokens', 'Every form includes a hidden token to prevent cross-site request forgery'],
        ['Rate Limiting', 'Login and upload endpoints are rate-limited to block automated attacks'],
        ['CSP Headers', 'Browser is instructed to block unauthorized scripts and resources'],
        ['Private GitHub Repo', 'Source code is not publicly visible — only the owner can access it'],
        ['Env Vars on Render', 'Secret keys are stored server-side, never in the code or GitHub'],
    ]
)
doc.add_paragraph()
divider()

# ════════════════════════════════════════════════
# 10. QUICK REFERENCE
# ════════════════════════════════════════════════
heading('10. Quick Reference', 1)

add_table(
    ['Task', 'How to Do It'],
    [
        ['Open the live app', 'https://filevault-eq7o.onrender.com'],
        ['View your code on GitHub', 'https://github.com/Dahirou-Bachar/FileVault'],
        ['View Render dashboard', 'https://dashboard.render.com'],
        ['See Render deploy logs', 'https://dashboard.render.com/web/srv-d81712gg4nts7398j0lg/logs'],
        ['Change environment variables', 'https://dashboard.render.com/web/srv-d81712gg4nts7398j0lg/env'],
        ['Push a code update', 'git add . → git commit -m "message" → git push'],
        ['Download code on new device', 'git clone https://github.com/Dahirou-Bachar/FileVault.git'],
        ['Run locally', 'venv\\Scripts\\activate → flask run'],
    ]
)

doc.add_paragraph()

# Footer note
foot_para = doc.add_paragraph()
foot_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
foot_run = foot_para.add_run(
    f'FileVault Project Documentation  •  Generated {datetime.date.today().strftime("%B %d, %Y")}'
)
foot_run.font.size = Pt(9)
foot_run.font.color.rgb = RGBColor(0x9c, 0xa3, 0xaf)
foot_run.font.italic = True

# Save
output_path = r'c:\Users\user\Desktop\FileVault_Project_Documentation.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
