from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ──
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2.5)

BLUE   = RGBColor(0x1a, 0x56, 0xdb)
DARK   = RGBColor(0x11, 0x18, 0x27)
GRAY   = RGBColor(0x6b, 0x72, 0x80)
GREEN  = RGBColor(0x05, 0x7A, 0x55)
RED    = RGBColor(0x9B, 0x1C, 0x1C)
ORANGE = RGBColor(0x92, 0x40, 0x09)

def heading(text, level=1, color=BLUE):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = color
        run.font.bold = True
    return p

def body(text, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = DARK
    return p

def bullet(text, bold_prefix=None, color=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.3)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.size = Pt(11)
        if color: r1.font.color.rgb = color
        r2 = p.add_run(text[len(bold_prefix):])
        r2.font.size = Pt(11)
    else:
        r = p.add_run(text)
        r.font.size = Pt(11)
        if color: r.font.color.rgb = color
    return p

def note(text, bg='FFF8E1', border='F59E0B'):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), bg)
    pPr.append(shd)
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    r.font.bold = True
    r.font.color.rgb = ORANGE
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.4)
    p.paragraph_format.right_indent = Inches(0.4)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F3F4F6')
    pPr.append(shd)
    r = p.add_run(text)
    r.font.name = 'Courier New'
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(0x1f, 0x29, 0x37)
    return p

def divider():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def table(headers, rows, header_color='1A56DB'):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), header_color)
        tcPr.append(shd)
    for ri, row_data in enumerate(rows):
        row = t.rows[ri + 1].cells
        for ci, txt in enumerate(row_data):
            row[ci].text = txt
            for para in row[ci].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
        if ri % 2 == 0:
            for ci in range(len(row_data)):
                tc = t.rows[ri + 1].cells[ci]._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'EFF6FF')
                tcPr.append(shd)
    return t

# ════════════════════════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════════════════════════
for _ in range(4):
    doc.add_paragraph()

tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run('FileVault')
r.font.size = Pt(42)
r.font.bold = True
r.font.color.rgb = BLUE

sp = doc.add_paragraph()
sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sp.add_run('Complete Project Documentation')
r.font.size = Pt(20)
r.font.color.rgb = GRAY

doc.add_paragraph()
dp = doc.add_paragraph()
dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = dp.add_run('Secure Cloud File Storage with Real-Time Monitoring')
r.font.size = Pt(13)
r.font.italic = True
r.font.color.rgb = DARK

for _ in range(2):
    doc.add_paragraph()

ip = doc.add_paragraph()
ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
lines = [
    f'Date: {datetime.date.today().strftime("%B %d, %Y")}',
    'Author: Dahirou Bachar',
    'Repository: github.com/Dahirou-Bachar/FileVault',
    'Live App: https://filevault-eq7o.onrender.com',
]
for line in lines:
    r = ip.add_run(line + '\n')
    r.font.size = Pt(11)
    r.font.color.rgb = GRAY

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ════════════════════════════════════════════════════════════
heading('Table of Contents', 1)
toc_items = [
    ('1.', 'Project Overview'),
    ('2.', 'What the Application Can Do — Full Feature List'),
    ('3.', 'Technology Stack'),
    ('4.', 'GitHub — Role in This Project'),
    ('5.', 'Render.com — Role in This Project'),
    ('6.', 'Data Persistence — What Is Saved and What Is Lost'),
    ('7.', 'Security Features'),
    ('8.', 'How to Update the Project'),
    ('9.', 'How to Recover on a New Device'),
    ('10.', 'Environment Variables Reference'),
    ('11.', 'Project File Structure'),
    ('12.', 'Quick Reference Links'),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    r1 = p.add_run(num + '  ')
    r1.font.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = BLUE
    r2 = p.add_run(title)
    r2.font.size = Pt(11)
doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 1. PROJECT OVERVIEW
# ════════════════════════════════════════════════════════════
heading('1. Project Overview')
body(
    'FileVault is a full-stack web application that provides secure, cloud-based file storage. '
    'It was built using Python (Flask) as the backend framework and Bootstrap 5 for the responsive '
    'frontend interface. The application is deployed live on the internet via Render.com and all '
    'source code is version-controlled on GitHub.'
)
doc.add_paragraph()
body(
    'The core purpose of FileVault is to allow registered users to upload, store, manage, and share '
    'files securely through a web browser — on any device including desktop computers, tablets, and '
    'mobile phones. Every file uploaded is encrypted before being saved, and users are protected by '
    'multiple layers of authentication security.'
)
doc.add_paragraph()
body('The application has three main layers:')
bullet('Frontend (what users see): HTML pages styled with Bootstrap 5 and custom CSS. Interactive features are powered by JavaScript and Bootstrap components.')
bullet('Backend (the engine): Python with Flask handles all logic — user login, file uploads, sharing, security checks, and database operations.')
bullet('Database (the memory): PostgreSQL stores all user accounts, passwords, file records, sharing history, and activity logs permanently.')
divider()

# ════════════════════════════════════════════════════════════
# 2. FEATURES
# ════════════════════════════════════════════════════════════
heading('2. What the Application Can Do — Full Feature List')

heading('2.1  User Registration & Login', 2)
bullet('New users register with their full name, email address, and a password')
bullet('Passwords are never stored in plain text — they are converted to a secure hash using bcrypt')
bullet('After registration, a verification email is sent — users must click the link to activate their account')
bullet('Login requires the correct email and password')
bullet('After 5 failed login attempts the account is locked for 15 minutes to prevent brute-force attacks')
bullet('Users can reset their password by requesting a link sent to their email')

heading('2.2  Two-Factor Authentication (2FA / TOTP)', 2)
bullet('Users can enable Two-Factor Authentication from their profile settings')
bullet('A QR code is displayed which the user scans with Google Authenticator or Authy')
bullet('After scanning, every login requires both the password AND a 6-digit code from the authenticator app')
bullet('The code changes every 30 seconds and is only valid once')
bullet('This means even if someone steals a password, they cannot log in without the phone')

heading('2.3  Biometric Login (WebAuthn / Passkey)', 2)
bullet('Users can register their fingerprint or Face ID as a login method')
bullet('Supported on phones with fingerprint sensors and laptops with Windows Hello or Touch ID')
bullet('Once registered, the user can log in by tapping their fingerprint instead of typing a password')
bullet('This uses the WebAuthn industry standard — the biometric data never leaves the device')

heading('2.4  File Upload', 2)
bullet('Users can upload files up to 50 MB each')
bullet('Supported file types: images (PNG, JPG, GIF, BMP, WEBP), documents (PDF, Word, Excel, PowerPoint), text files (TXT, CSV, JSON, XML), archives (ZIP, RAR), and media (MP4, MP3)')
bullet('Files can be uploaded by clicking the Upload button or by dragging and dropping onto the upload area')
bullet('A real-time progress bar shows upload percentage')
bullet('Each user has a storage quota of 500 MB')
bullet('Every file is encrypted using Fernet symmetric encryption before being saved to disk')

heading('2.5  File Management', 2)
bullet('Users see all their files displayed as cards with the file type icon or image preview')
bullet('Files can be sorted by date, name, or size — in ascending or descending order')
bullet('Files can be filtered by type: images, documents, spreadsheets, media, or archives')
bullet('Files can be searched by name using the search bar')
bullet('Each file shows its size and upload date')
bullet('Files can be renamed without re-uploading')
bullet('Files can be permanently deleted')

heading('2.6  Decrypt & Download', 2)
bullet('Every file has a Download button (blue arrow icon) — this downloads the file directly')
bullet('Every encrypted file also has a Decrypt & Download button (yellow unlock icon)')
bullet('Clicking the unlock button opens a prompt asking for the user\'s account password')
bullet('If the correct password is entered, the file is decrypted on the server and downloaded')
bullet('If the wrong password is entered, an error message is shown and nothing is downloaded')
bullet('This adds an extra layer of protection — even if someone accesses the account, they need the password to get the file')

heading('2.7  File Sharing', 2)
bullet('Owners can share files with other registered users of the application')
bullet('Shared users can view and download the file from their "Shared with me" section')
bullet('Owners can also generate a public share link with an optional expiry date (1 day, 7 days, 30 days, or never)')
bullet('Anyone with the public link can download the file — no account needed')
bullet('Public links can be revoked at any time by the owner')

heading('2.8  Real-Time Notifications', 2)
bullet('The dashboard shows live activity notifications using WebSocket technology (Socket.IO)')
bullet('When a file is uploaded, deleted, or shared, a notification bar appears at the top of the dashboard')
bullet('Notifications from other sessions (e.g. another device) trigger an automatic page refresh')
bullet('Notifications auto-dismiss after 6 seconds')

heading('2.9  Activity Log', 2)
bullet('Every action in the application is recorded: uploads, downloads, deletions, shares, logins, logouts, registrations, and profile changes')
bullet('Users can view their full activity history on the Activity page')
bullet('Each log entry shows the action type, description, and timestamp')
bullet('Admins can view all users\' activity')

heading('2.10  Admin Panel', 2)
bullet('Admin users have access to a special Admin Panel page')
bullet('Admins can see all registered users and their file counts')
bullet('Admins can manage users and monitor the platform')
bullet('The first registered account with is_admin=True has full admin access')

heading('2.11  User Profile & Settings', 2)
bullet('Users can update their full name and email address')
bullet('Users can change their password')
bullet('Users can enable or disable Two-Factor Authentication')
bullet('Users can register or remove biometric login')
bullet('The profile page shows account statistics (files uploaded, storage used)')

heading('2.12  Dark Mode', 2)
bullet('Users can switch between light mode and dark mode using the theme toggle button')
bullet('The chosen theme is saved in the browser — it persists across page refreshes and visits')

heading('2.13  Responsive Design (Mobile Friendly)', 2)
bullet('The interface works on all screen sizes: desktop, tablet, and mobile phone')
bullet('On small screens the sidebar collapses — it can be opened with the 3-bar hamburger button')
bullet('Pressing Escape on the keyboard also closes the sidebar')
bullet('Tapping the dark overlay behind the sidebar closes it')
divider()

# ════════════════════════════════════════════════════════════
# 3. TECHNOLOGY STACK
# ════════════════════════════════════════════════════════════
heading('3. Technology Stack')
body('The following technologies were used to build FileVault:')
doc.add_paragraph()
table(
    ['Layer', 'Technology', 'Version', 'Purpose'],
    [
        ['Backend', 'Python', '3.11', 'Core programming language'],
        ['Web Framework', 'Flask', '3.0.3', 'Handles HTTP requests, routing, and application logic'],
        ['Database ORM', 'SQLAlchemy + Flask-Migrate', '2.0 / 4.0', 'Manages database models and schema migrations'],
        ['Production Database', 'PostgreSQL', 'Latest', 'Permanent storage of users, files metadata, activity logs'],
        ['Local Dev Database', 'SQLite', 'Built-in', 'Lightweight database for local development only'],
        ['Authentication', 'Flask-Login + Werkzeug', '0.6 / 3.0', 'Session management and secure password hashing (bcrypt)'],
        ['Two-Factor Auth', 'PyOTP + QRCode', '2.9 / 8.0', 'Generates TOTP codes and QR codes for authenticator apps'],
        ['Biometric Auth', 'WebAuthn', '2.2.0', 'Fingerprint and Face ID login using FIDO2 standard'],
        ['File Encryption', 'Cryptography (Fernet)', '43.0', 'Symmetric AES encryption of uploaded files at rest'],
        ['Real-Time', 'Flask-SocketIO + Eventlet', '5.3 / 0.36', 'WebSocket connections for live dashboard notifications'],
        ['Email', 'Flask-Mail', '0.10', 'Sends verification emails and password reset links via Gmail'],
        ['Rate Limiting', 'Flask-Limiter', '4.1', 'Prevents brute-force attacks by limiting request rates'],
        ['CSRF Protection', 'Flask-WTF', '1.2', 'Protects all forms against cross-site request forgery'],
        ['Frontend Framework', 'Bootstrap 5', '5.3.3', 'Responsive layout, components, and utility classes'],
        ['Icons', 'Bootstrap Icons', '1.11.3', 'Over 1,800 SVG icons used throughout the interface'],
        ['Charts', 'Chart.js', '4.4.3', 'Renders activity charts on the dashboard'],
        ['WSGI Server', 'Gunicorn + Eventlet worker', '22.0', 'Production web server that handles concurrent requests'],
        ['Version Control', 'Git + GitHub', 'Latest', 'Code storage, history, and deployment trigger'],
        ['Cloud Hosting', 'Render.com', 'Free tier', 'Runs the live application 24/7 on the internet'],
    ]
)
doc.add_paragraph()
divider()

# ════════════════════════════════════════════════════════════
# 4. GITHUB
# ════════════════════════════════════════════════════════════
heading('4. GitHub — Role in This Project')

heading('4.1  What is GitHub?', 2)
body(
    'GitHub is a cloud platform that stores code using the Git version control system. '
    'It keeps a full history of every change ever made to the project — who changed what, '
    'when, and with a description of why. It also acts as the trigger that tells Render '
    'to update the live application whenever new code is pushed.'
)

heading('4.2  How GitHub is Used in FileVault', 2)
bullet('Code Backup:', 'Code Backup:')
body('     All project files — Python code, HTML templates, CSS styles, and JavaScript — are stored in a private repository called FileVault under the account Dahirou-Bachar. Private means only you can see and access it.')
doc.add_paragraph()
bullet('Version History:', 'Version History:')
body('     Every time changes are saved (committed), GitHub records the full change. You can always go back to any previous version of the code.')
doc.add_paragraph()
bullet('Deployment Trigger:', 'Deployment Trigger:')
body('     Render.com watches the GitHub repository. The moment new code is pushed to GitHub, Render automatically downloads it and redeploys the live application — usually within 2 to 3 minutes.')
doc.add_paragraph()
bullet('Disaster Recovery:', 'Disaster Recovery:')
body('     If the computer is lost or broken, the entire project can be re-downloaded from GitHub to any new device with one command.')

heading('4.3  Repository Information', 2)
table(
    ['Property', 'Value'],
    [
        ['Repository Name', 'FileVault'],
        ['Owner Account', 'Dahirou-Bachar'],
        ['Visibility', 'Private — only the owner can see and clone it'],
        ['Main Branch', 'main'],
        ['GitHub URL', 'https://github.com/Dahirou-Bachar/FileVault'],
    ]
)
doc.add_paragraph()

heading('4.4  Git Commands Used to Push Updates', 2)
body('When code changes are made locally, these three commands send them to GitHub:')
doc.add_paragraph()
code_block('git add .')
body('     Marks all changed files as ready to be saved.')
doc.add_paragraph()
code_block('git commit -m "Short description of what changed"')
body('     Saves the changes locally with a label.')
doc.add_paragraph()
code_block('git push')
body('     Sends the saved changes to GitHub, which triggers Render to redeploy.')
divider()

# ════════════════════════════════════════════════════════════
# 5. RENDER
# ════════════════════════════════════════════════════════════
heading('5. Render.com — Role in This Project')

heading('5.1  What is Render?', 2)
body(
    'Render.com is a cloud hosting platform that runs the FileVault application on a server '
    '24 hours a day, 7 days a week. Without a hosting service like Render, the app would only '
    'work when your personal computer is turned on and running Flask locally.'
)

heading('5.2  What Render Does for FileVault', 2)
bullet('Runs the Application:', 'Runs the Application:')
body('     Render executes the Gunicorn web server which loads the Flask application and listens for visitor requests on a public URL.')
doc.add_paragraph()
bullet('Auto-Deploy from GitHub:', 'Auto-Deploy from GitHub:')
body('     Every time code is pushed to GitHub, Render detects the change, pulls the new code, installs dependencies, and restarts the application automatically.')
doc.add_paragraph()
bullet('Hosts the PostgreSQL Database:', 'Hosts the PostgreSQL Database:')
body('     Render provides a free PostgreSQL database service where all user accounts, passwords, file records, and activity logs are stored permanently — separate from the application server.')
doc.add_paragraph()
bullet('Stores Environment Variables Securely:', 'Stores Environment Variables Securely:')
body('     Secret keys (encryption key, secret key, email password) are stored in Render\'s environment settings — never in the code or on GitHub.')
doc.add_paragraph()
bullet('Provides a Public URL:', 'Provides a Public URL:')
body('     The application is accessible from anywhere in the world at: https://filevault-eq7o.onrender.com')

heading('5.3  Render Service Details', 2)
table(
    ['Property', 'Value'],
    [
        ['Service Name', 'FileVault'],
        ['Service Type', 'Web Service'],
        ['Plan', 'Free Tier'],
        ['Live URL', 'https://filevault-eq7o.onrender.com'],
        ['Region', 'Oregon, USA'],
        ['Python Version', '3.11'],
        ['Start Command', 'flask db upgrade && gunicorn --worker-class eventlet -w 1 --bind 0.0.0.0:$PORT app:app'],
        ['Auto-Deploy', 'Yes — triggered automatically on every GitHub push'],
        ['Database', 'PostgreSQL (free tier, persistent)'],
    ]
)
doc.add_paragraph()

heading('5.4  Free Tier Limitation — Sleep Mode', 2)
body(
    'On Render\'s free plan, the application server "sleeps" after 15 minutes of no visitors '
    'to save computing resources. The first request after sleeping takes 30 to 60 seconds to '
    'wake up — this delay is completely normal. Once awake, the app runs at full speed. '
    'Upgrading to Render\'s paid plan ($7/month) removes the sleep behaviour entirely.'
)
divider()

# ════════════════════════════════════════════════════════════
# 6. DATA PERSISTENCE
# ════════════════════════════════════════════════════════════
heading('6. Data Persistence — What Is Saved and What Is Lost')
body(
    'Understanding what survives a redeploy is important. When Render redeploys the application '
    '(triggered by a GitHub push), two different storage systems are involved:'
)
doc.add_paragraph()

heading('6.1  What is Stored Where', 2)
table(
    ['Data', 'Stored In', 'Survives Redeploy?'],
    [
        ['User accounts (email, full name)', 'PostgreSQL Database', 'YES — forever'],
        ['Passwords (hashed)', 'PostgreSQL Database', 'YES — forever'],
        ['File names and metadata', 'PostgreSQL Database', 'YES — forever'],
        ['Sharing records', 'PostgreSQL Database', 'YES — forever'],
        ['Activity logs', 'PostgreSQL Database', 'YES — forever'],
        ['Actual file content (the uploaded files)', 'Render Server Disk', 'NO — wiped on redeploy'],
    ]
)
doc.add_paragraph()
note('⚠  Important: Uploaded file content is stored on Render\'s temporary disk and is wiped every time the app redeploys. User accounts and all database records are permanent.')

heading('6.2  Why File Content Is Lost', 2)
body(
    'Render\'s free plan uses "ephemeral" (temporary) disk storage. Each time the application is '
    'redeployed, Render creates a fresh server environment. Any files written to the disk during '
    'the previous deployment — including the vault_uploads folder — are gone.'
)

heading('6.3  How to Make File Content Permanent (Future Upgrade)', 2)
body(
    'To make uploaded files survive redeploys, they must be moved to a cloud object storage '
    'service. The recommended option for this project is Cloudflare R2, which offers:'
)
bullet('10 GB of free storage per month')
bullet('10 million free download operations per month')
bullet('No data transfer fees')
bullet('Compatible with the same boto3 library used for AWS S3')
body('This would require updating the upload, download, and delete routes in app.py to communicate with the R2 bucket instead of the local disk. This is a planned future improvement.')
divider()

# ════════════════════════════════════════════════════════════
# 7. SECURITY
# ════════════════════════════════════════════════════════════
heading('7. Security Features')
body('FileVault was built with security as a primary concern. The following protections are in place:')
doc.add_paragraph()
table(
    ['Security Feature', 'Description', 'Protects Against'],
    [
        ['Bcrypt Password Hashing', 'Passwords are converted to an irreversible hash before storage', 'Database leaks — the real password cannot be recovered from the hash'],
        ['File Encryption (Fernet AES)', 'Every uploaded file is encrypted before being written to disk', 'Server access — files cannot be read without the encryption key'],
        ['Decrypt with Password', 'Users must re-enter their password to download decrypted files', 'Unauthorized access to files even if someone is logged in'],
        ['Two-Factor Authentication (TOTP)', 'A time-based one-time code is required at login', 'Stolen passwords — attacker cannot log in without the phone'],
        ['WebAuthn / Biometrics', 'Fingerprint or Face ID as an alternative login method', 'Password theft — biometric data never leaves the device'],
        ['CSRF Tokens', 'Every form includes a hidden verification token', 'Cross-site request forgery attacks from malicious websites'],
        ['Rate Limiting', 'Login and upload requests are limited per IP address', 'Automated brute-force and flooding attacks'],
        ['Account Lockout', 'Account locked for 15 minutes after 5 failed logins', 'Password guessing attacks'],
        ['Email Verification', 'New accounts must verify their email before full access', 'Fake account creation with invalid emails'],
        ['Content Security Policy', 'Browser instructed to block unauthorized external scripts', 'Cross-site scripting (XSS) attacks'],
        ['HTTPS Only (on Render)', 'All traffic is encrypted in transit via TLS/SSL', 'Network eavesdropping and man-in-the-middle attacks'],
        ['Private GitHub Repository', 'Source code is not publicly visible', 'Attackers reading the code to find vulnerabilities'],
        ['Environment Variables', 'Secret keys stored on Render, not in the code', 'Secret leakage through GitHub if code is ever made public'],
    ]
)
doc.add_paragraph()
divider()

# ════════════════════════════════════════════════════════════
# 8. HOW TO UPDATE
# ════════════════════════════════════════════════════════════
heading('8. How to Update the Project')
body(
    'Because GitHub and Render are connected, updating the live application is simple. '
    'Changes are made on the local computer, saved to GitHub, and Render handles the rest automatically.'
)
doc.add_paragraph()

heading('8.1  The Update Process', 2)
table(
    ['Step', 'Action', 'Time'],
    [
        ['1', 'Make changes to the code using VS Code or any editor', 'As long as needed'],
        ['2', 'Run: git add .', 'Instant'],
        ['3', 'Run: git commit -m "Description of the change"', 'Instant'],
        ['4', 'Run: git push', 'A few seconds'],
        ['5', 'Render detects the push and starts rebuilding automatically', '2 to 3 minutes'],
        ['6', 'Live app is updated — users see the new version', 'Immediate after build'],
    ]
)
doc.add_paragraph()

heading('8.2  What Happens During Render Rebuild', 2)
bullet('Render pulls the latest code from GitHub')
bullet('It installs all Python packages listed in requirements.txt')
bullet('It runs flask db upgrade to apply any database changes')
bullet('It starts Gunicorn with the new code')
bullet('The live URL is now serving the updated version')

heading('8.3  Monitoring the Deployment', 2)
body('The deployment progress can be watched in real time at:')
code_block('https://dashboard.render.com/web/srv-d81712gg4nts7398j0lg/logs')
body('The log shows every step of the build and any errors that occur.')
divider()

# ════════════════════════════════════════════════════════════
# 9. NEW DEVICE RECOVERY
# ════════════════════════════════════════════════════════════
heading('9. How to Recover on a New Device')
body(
    'If the computer is lost, reset, or replaced, the entire project can be recovered because '
    'all code is stored on GitHub and the live app keeps running on Render regardless.'
)
doc.add_paragraph()

heading('Step 1 — Install Required Software', 2)
bullet('Python 3.11 — download from python.org')
bullet('Git — download from git-scm.com')
bullet('VS Code (optional) — download from code.visualstudio.com')

heading('Step 2 — Download the Code from GitHub', 2)
code_block('git clone https://github.com/Dahirou-Bachar/FileVault.git\ncd FileVault')

heading('Step 3 — Set Up Python Environment', 2)
code_block('python -m venv venv\nvenv\\Scripts\\activate\npip install -r requirements.txt')

heading('Step 4 — Restore the Secret Keys', 2)
body('Go to Render\'s environment variables page and copy each variable value:')
code_block('https://dashboard.render.com/web/srv-d81712gg4nts7398j0lg/env')
body('Create a file named .env in the project folder and paste the values inside.')

heading('Step 5 — Run Locally', 2)
code_block('flask db upgrade\nflask run')
body('The app will be available at http://localhost:5000')

note('Note: The live app on Render continues to run and is unaffected by any of this. Recovery only affects local development.')
divider()

# ════════════════════════════════════════════════════════════
# 10. ENVIRONMENT VARIABLES
# ════════════════════════════════════════════════════════════
heading('10. Environment Variables Reference')
body(
    'Environment variables are configuration values that are kept secret and stored on Render\'s '
    'servers — never inside the code or on GitHub. The application reads them automatically at startup.'
)
doc.add_paragraph()
table(
    ['Variable Name', 'Purpose', 'Required?'],
    [
        ['SECRET_KEY', 'Encrypts user sessions and cookies. Must be a long random string.', 'Yes'],
        ['ENCRYPTION_KEY', 'Fernet key used to encrypt and decrypt uploaded files. Must be a valid 32-byte base64 key.', 'Yes'],
        ['DATABASE_URL', 'PostgreSQL connection string provided by Render. Tells the app where the database is.', 'Yes (for persistence)'],
        ['MAIL_USERNAME', 'Gmail address used to send verification and password-reset emails.', 'Yes (for email features)'],
        ['MAIL_PASSWORD', 'Gmail App Password (16-character code from Google account settings — not the regular Gmail password).', 'Yes (for email features)'],
        ['WEBAUTHN_RP_ID', 'Domain name of the live site (e.g. filevault-eq7o.onrender.com). Required for biometric login.', 'Yes (for biometrics)'],
        ['WEBAUTHN_RP_ORIGIN', 'Full HTTPS URL of the live site (e.g. https://filevault-eq7o.onrender.com). Must match exactly.', 'Yes (for biometrics)'],
        ['PYTHON_VERSION', 'Forces Render to use Python 3.11. Set to: 3.11.0', 'Yes'],
    ]
)
doc.add_paragraph()
divider()

# ════════════════════════════════════════════════════════════
# 11. FILE STRUCTURE
# ════════════════════════════════════════════════════════════
heading('11. Project File Structure')
body('The project folder is organised as follows:')
doc.add_paragraph()
table(
    ['File / Folder', 'Purpose'],
    [
        ['app.py', 'The main application file — contains all routes, models, and business logic (over 1,400 lines)'],
        ['requirements.txt', 'Lists every Python package the project depends on. Render uses this to install dependencies.'],
        ['.python-version', 'Tells Render and local tools to use Python 3.11 specifically'],
        ['.gitignore', 'Lists files that must NOT be pushed to GitHub (secret .env file, database, uploaded files, cache)'],
        ['templates/', 'Folder containing all HTML page templates'],
        ['templates/base.html', 'The master layout shared by all pages (header, footer, scripts)'],
        ['templates/dashboard.html', 'The main file management page — shows file cards, upload modal, sharing'],
        ['templates/auth/', 'Login, register, password reset, 2FA setup and verification pages'],
        ['templates/profile.html', 'User profile and settings page'],
        ['templates/activity.html', 'Full activity log page'],
        ['templates/admin.html', 'Admin panel for managing all users'],
        ['static/css/style.css', 'Custom CSS styles for the entire application (dark mode, sidebar, cards, etc.)'],
        ['static/js/main.js', 'Global JavaScript — dark mode toggle, sidebar toggle, toast notifications'],
        ['migrations/', 'Database migration scripts generated by Flask-Migrate. Tracks schema changes over time.'],
        ['instance/vault.db', 'SQLite database file used for local development only (excluded from GitHub)'],
        ['vault_uploads/', 'Folder where encrypted uploaded files are stored locally (excluded from GitHub)'],
        ['.env', 'Local secret configuration file (excluded from GitHub — never commit this file)'],
    ]
)
doc.add_paragraph()
divider()

# ════════════════════════════════════════════════════════════
# 12. QUICK REFERENCE
# ════════════════════════════════════════════════════════════
heading('12. Quick Reference Links')
table(
    ['Task', 'Link / Command'],
    [
        ['Open the live application', 'https://filevault-eq7o.onrender.com'],
        ['View code on GitHub', 'https://github.com/Dahirou-Bachar/FileVault'],
        ['Render dashboard', 'https://dashboard.render.com'],
        ['Render deploy logs', 'https://dashboard.render.com/web/srv-d81712gg4nts7398j0lg/logs'],
        ['Render environment variables', 'https://dashboard.render.com/web/srv-d81712gg4nts7398j0lg/env'],
        ['Render service settings', 'https://dashboard.render.com/web/srv-d81712gg4nts7398j0lg/settings'],
        ['Push a code update', 'git add .   →   git commit -m "message"   →   git push'],
        ['Clone to new device', 'git clone https://github.com/Dahirou-Bachar/FileVault.git'],
        ['Run locally', 'venv\\Scripts\\activate   →   flask run'],
        ['View deploy history on GitHub', 'https://github.com/Dahirou-Bachar/FileVault/commits/main'],
    ]
)
doc.add_paragraph()

# Footer
fp = doc.add_paragraph()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run(
    f'FileVault — Complete Project Documentation  •  {datetime.date.today().strftime("%B %d, %Y")}  •  Dahirou Bachar'
)
fr.font.size = Pt(9)
fr.font.color.rgb = GRAY
fr.font.italic = True

out = r'c:\Users\user\Desktop\FileVault_Complete_Documentation.docx'
doc.save(out)
print(f'Saved: {out}')
